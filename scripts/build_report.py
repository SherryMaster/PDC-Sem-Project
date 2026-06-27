#!/usr/bin/env python3
"""Generate report.docx from report_content.json + figures."""
import json
import sys
from pathlib import Path

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Inches, Pt

SCRIPT_DIR = Path(__file__).resolve().parent
PROJECT_ROOT = SCRIPT_DIR.parent
CONTENT_PATH = SCRIPT_DIR / "report_content.json"
LOGO_PATH = SCRIPT_DIR / "figures" / "ucp_logo.jpg"
OUTPUT_PATH = PROJECT_ROOT / "report.docx"


def load_content() -> dict:
    with open(CONTENT_PATH) as f:
        return json.load(f)


def add_heading(doc: Document, text: str, level: int = 1) -> None:
    h = doc.add_heading(text, level=level)
    h.alignment = WD_ALIGN_PARAGRAPH.LEFT


def add_body(doc: Document, text: str) -> None:
    p = doc.add_paragraph(text)
    p.paragraph_format.space_after = Pt(8)


def build_cover(doc: Document, content: dict) -> None:
    cover = content["cover"]

    if LOGO_PATH.exists():
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = p.add_run()
        run.add_picture(str(LOGO_PATH), width=Inches(2.0))

    title_blocks = [
        cover["university"],
        cover["program"],
        cover["project_type"],
    ]
    for line in title_blocks:
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = p.add_run(line)
        run.bold = True
        run.font.size = Pt(20 if line == cover["university"] else 16)

    doc.add_paragraph()
    topic_p = doc.add_paragraph()
    topic_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    topic_run = topic_p.add_run("Topic: " + cover["topic"])
    topic_run.bold = True
    topic_run.font.size = Pt(14)

    doc.add_paragraph()
    doc.add_paragraph()

    info_lines = [
        f"Submitted by: {cover['submitted_by_name']}",
        f"Registration Number: {cover['submitted_by_reg']}",
        f"Degree Program: {cover['submitted_by_degree']}",
        f"Tools Used: {cover['tools']}",
    ]
    for line in info_lines:
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.LEFT
        run = p.add_run(line)
        run.font.size = Pt(13)

    doc.add_page_break()


def build_introduction(doc: Document, content: dict) -> None:
    add_heading(doc, "1. Introduction", level=1)
    add_body(doc, content["introduction"])


def build_selected_modules(doc: Document, content: dict) -> None:
    add_heading(doc, "2. Selected Functional Modules", level=1)
    add_body(doc, content["selected_modules_intro"])
    blurbs = content["module_blurbs"]
    add_heading(doc, "2.1 Sequential Frame Analyzer", level=2)
    add_body(doc, blurbs["sequential"])
    add_heading(doc, "2.2 Pthread Frame Analyzer", level=2)
    add_body(doc, blurbs["pthread"])
    add_heading(doc, "2.3 OpenMP Frame Analyzer", level=2)
    add_body(doc, blurbs["openmp"])


def build_logic_explanation(doc: Document, content: dict) -> None:
    add_heading(doc, "3. Original Logic Explanation", level=1)
    logic = content["logic_explanations"]
    add_heading(doc, "3.1 Sequential Logic", level=2)
    add_body(doc, logic["sequential"])
    add_heading(doc, "3.2 Pthread Logic", level=2)
    add_body(doc, logic["pthread"])
    add_heading(doc, "3.3 OpenMP Logic", level=2)
    add_body(doc, logic["openmp"])


def add_code_listing(doc: Document, title: str, code: str, source_file: str, source_lines: str) -> None:
    p_title = doc.add_paragraph()
    title_run = p_title.add_run(title)
    title_run.bold = True
    title_run.font.size = Pt(11)
    for line in code.split("\n"):
        p = doc.add_paragraph()
        p.paragraph_format.left_indent = Inches(0.25)
        p.paragraph_format.space_after = Pt(0)
        run = p.add_run(line if line else " ")
        run.font.name = "Courier New"
        run.font.size = Pt(9)
    p_cap = doc.add_paragraph()
    cap_run = p_cap.add_run(f"Source: {source_file}, lines {source_lines}")
    cap_run.italic = True
    cap_run.font.size = Pt(9)
    p_cap.paragraph_format.space_after = Pt(8)


def build_c_implementations(doc: Document, content: dict) -> None:
    add_heading(doc, "4. C Implementations", level=1)
    add_body(doc, "Each subsection below shows the kernel of one execution mode. The shared Sobel 3x3 convolution is shown in 4.0; the three mode-specific branches of the main loop are shown in 4.1 (Sequential), 4.2 (Pthread), and 4.3 (OpenMP).")
    listings = content["code_listings"]
    add_heading(doc, "4.0 Sobel Convolution (shared kernel)", level=2)
    add_code_listing(doc, listings["sobel"]["title"], listings["sobel"]["code"], listings["sobel"]["source_file"], listings["sobel"]["source_lines"])
    add_heading(doc, "4.1 Sequential Frame Analyzer", level=2)
    add_body(doc, "The Sequential branch is a plain for-loop over the frame list:")
    p = doc.add_paragraph()
    p.paragraph_format.left_indent = Inches(0.25)
    run = p.add_run('for (int i = 0; i < file_count; i++) {\n    long long frame_edges = 0;\n    process_single_frame(jobs[i].filepath, output_dir, &frame_edges);\n    total_edges += frame_edges;\n}')
    run.font.name = "Courier New"
    run.font.size = Pt(9)
    add_heading(doc, "4.2 Pthread Frame Analyzer", level=2)
    add_code_listing(doc, listings["pthread"]["title"], listings["pthread"]["code"], listings["pthread"]["source_file"], listings["pthread"]["source_lines"])
    add_heading(doc, "4.3 OpenMP Frame Analyzer", level=2)
    add_code_listing(doc, listings["openmp"]["title"], listings["openmp"]["code"], listings["openmp"]["source_file"], listings["openmp"]["source_lines"])


def build_gui_section(doc: Document, content: dict) -> None:
    add_heading(doc, "5. GUI Frontend (Tkinter)", level=1)
    add_body(doc, "The Parallel Batch Frame Analyzer ships with a Python Tkinter GUI that wraps the C binary. The GUI provides three tabs: Scanner (run a single mode with optional worker count), Scan History (a session-scoped log of completed runs), and Benchmark Report (run all three modes back-to-back with side-by-side comparison and export to .txt or .csv). Below is the FrameAnalyzerApp class skeleton showing how mode, worker count, and the input directory are wired into the underlying subprocess calls.")
    add_code_listing(
        doc,
        content["code_listings"]["gui"]["title"],
        content["code_listings"]["gui"]["code"],
        content["code_listings"]["gui"]["source_file"],
        content["code_listings"]["gui"]["source_lines"],
    )


def build_system_specs(doc: Document, content: dict) -> None:
    add_heading(doc, "6. System Specifications", level=1)
    s = content["system_specs"]
    lines = [
        f"CPU Model       : {s['cpu_model']}",
        f"Physical Cores  : {s['physical_cores']}",
        f"Logical Threads : {s['logical_threads']}",
        f"Total Memory    : {s['total_memory_gb']} GB",
        f"Kernel Version  : {s['kernel_version']}",
        f"Compiler        : {s['compiler']}",
        f"OpenMP Threads  : {s['openmp_threads']}",
        f"Python Version  : {s['python_version']}",
        f"Operating System: {s['os']}",
    ]
    for line in lines:
        p = doc.add_paragraph(line)
        p.paragraph_format.left_indent = Inches(0.25)
        for run in p.runs:
            run.font.name = "Courier New"
            run.font.size = Pt(10)


def build_build_and_execution(doc: Document, content: dict) -> None:
    add_heading(doc, "7. Build & Execution", level=1)
    b = content["build_execution"]
    add_body(doc, "The project is built with a single `make` invocation. All source files are compiled and linked into a single binary.")
    for label, value in [
        ("Compiler", b["compiler"]),
        ("Linker flags", b["linker_flags"]),
        ("Build command", b["build_command"]),
        ("Output binary", b["binary"]),
    ]:
        p = doc.add_paragraph()
        run_label = p.add_run(f"{label}: ")
        run_label.bold = True
        p.add_run(value)
    add_heading(doc, "7.1 CLI Usage", level=2)
    for cmd in b["cli_modes"]:
        p = doc.add_paragraph(cmd)
        p.paragraph_format.left_indent = Inches(0.25)
        for run in p.runs:
            run.font.name = "Courier New"
            run.font.size = Pt(10)
    add_heading(doc, "7.2 GUI Launch", level=2)
    doc.add_paragraph(b["gui_command"])
    add_heading(doc, "7.3 Test Frame Generation", level=2)
    doc.add_paragraph(b["generate_command"])


def build_test_workload(doc: Document, content: dict) -> None:
    add_heading(doc, "8. Test Workload & Figures", level=1)
    tw = content["test_workload"]
    add_body(doc, tw["description"])
    add_heading(doc, "8.1 Sample Figures", level=2)
    figures = tw["figures"]
    if len(figures) >= 2:
        table = doc.add_table(rows=1, cols=2)
        table.alignment = WD_ALIGN_PARAGRAPH.CENTER
        for col_idx, fig in enumerate(figures):
            cell = table.rows[0].cells[col_idx]
            cell_p = cell.paragraphs[0]
            cell_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            img_path = PROJECT_ROOT / fig["path"]
            if img_path.exists():
                cell_p.add_run().add_picture(str(img_path), width=Inches(3.0))
            cap_p = cell.add_paragraph(fig["caption"])
            cap_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            for run in cap_p.runs:
                run.italic = True
                run.font.size = Pt(9)


def build_timing_comparison(doc: Document, content: dict) -> None:
    add_heading(doc, "9. Completion Time Comparison", level=1)
    t = content["timing_table"]
    add_body(doc, f"All three modes produce exactly {t['edge_pixels']:,} edge pixels across the 20-frame test workload, confirming that the parallelization preserves correctness. The table below reports wall-clock completion time in milliseconds for each mode at worker counts 1, 2, 4, and 8.")
    table = doc.add_table(rows=1 + len(t["rows"]), cols=len(t["headers"]))
    table.style = "Light Grid Accent 1"
    for col_idx, header in enumerate(t["headers"]):
        cell = table.rows[0].cells[col_idx]
        cell_p = cell.paragraphs[0]
        run = cell_p.add_run(header)
        run.bold = True
    for row_idx, row in enumerate(t["rows"], start=1):
        for col_idx, value in enumerate(row):
            cell = table.rows[row_idx].cells[col_idx]
            cell_p = cell.paragraphs[0]
            cell_p.add_run(value)
    add_heading(doc, "9.1 Analysis", level=2)
    add_body(doc, t["analysis"])


def build_conclusion(doc: Document, content: dict) -> None:
    add_heading(doc, "10. Conclusion", level=1)
    add_body(doc, content["conclusion"])


def build_references(doc: Document, content: dict) -> None:
    add_heading(doc, "11. References", level=1)
    add_body(doc, "American Psychological Association (APA) format.")
    for ref in content["references"]:
        p = doc.add_paragraph(ref)
        p.paragraph_format.left_indent = Inches(0.5)
        p.paragraph_format.first_line_indent = Inches(-0.5)
        for run in p.runs:
            run.font.size = Pt(11)


def main() -> int:
    if not CONTENT_PATH.exists():
        print(f"ERROR: missing {CONTENT_PATH}", file=sys.stderr)
        return 1

    content = load_content()
    doc = Document()

    build_cover(doc, content)
    build_introduction(doc, content)
    build_selected_modules(doc, content)
    build_logic_explanation(doc, content)
    build_c_implementations(doc, content)
    build_gui_section(doc, content)
    build_system_specs(doc, content)
    build_build_and_execution(doc, content)
    build_test_workload(doc, content)
    build_timing_comparison(doc, content)
    build_conclusion(doc, content)
    build_references(doc, content)

    doc.save(str(OUTPUT_PATH))
    print(f"Wrote {OUTPUT_PATH}")
    return 0


if __name__ == "__main__":
    sys.exit(main())
